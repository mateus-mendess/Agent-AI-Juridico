import os
import logging
import tempfile
from openai import OpenAI
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

client = OpenAI(
    api_key=os.getenv("OPENROUTER_API_KEY"),
    base_url="https://openrouter.ai/api/v1"
)

MODEL = "openrouter/auto"

TIPOS_DOCUMENTO = {
    "doc_procuracao": "Procuração",
    "doc_termo": "Termo de Declaração",
    "doc_contrato": "Contrato",
}

CAMPOS_DOCUMENTO = {
    "doc_procuracao": [
        ("outorgante_nome", "Nome completo do *Outorgante*:"),
        ("outorgante_cpf", "CPF do Outorgante:"),
        ("outorgante_qualificacao", "Qualificação do Outorgante (ex: brasileiro, policial militar, solteiro):"),
        ("outorgante_endereco", "Endereço completo do Outorgante:"),
        ("outorgado_nome", "Nome completo do *Outorgado* (advogado):"),
        ("outorgado_oab", "Número da OAB (ex: OAB/AL nº 15.787):"),
        ("outorgado_endereco", "Endereço profissional do Outorgado:"),
        ("tipo", "Tipo da procuração (*cível* ou *criminal*):"),
        ("numero_processo", "Número do processo (apenas se criminal, senão digite *não se aplica*):"),
        ("cidade", "Cidade e Estado:"),
    ],
    "doc_termo": [
        ("declarante_nome", "Nome completo do *Declarante*:"),
        ("declarante_qualificacao", "Qualificação (ex: brasileiro, policial militar, solteiro):"),
        ("declarante_cpf", "CPF do Declarante:"),
        ("declarante_endereco", "Endereço completo do Declarante:"),
        ("cidade", "Cidade e Estado:"),
    ],
    "doc_contrato": [
        ("advogado_nome", "Nome completo do *Advogado*:"),
        ("advogado_oab", "Número da OAB (ex: OAB/AL nº 15.787):"),
        ("advogado_endereco", "Endereço profissional do Advogado:"),
        ("cliente_nome", "Nome completo do *Cliente*:"),
        ("cliente_qualificacao", "Qualificação do Cliente (ex: brasileiro, casado, fotógrafo):"),
        ("cliente_cpf", "CPF do Cliente:"),
        ("cliente_endereco", "Endereço completo do Cliente:"),
        ("valor_entrada", "Valor de entrada (ex: R$ 500,00 (quinhentos reais)):"),
        ("percentual", "Percentual sobre o proveito (ex: 30%):"),
        ("cidade", "Cidade e Estado:"),
    ],
}

PROMPTS_DOCUMENTO = {
    "doc_procuracao": """Gere uma Procuração formal em português brasileiro seguindo EXATAMENTE esta estrutura, sem adicionar nada além do que está aqui:

OUTORGANTE(s): {outorgante_nome}, {outorgante_qualificacao}, inscrito no CPF nº {outorgante_cpf}, residente e domiciliado em {outorgante_endereco}.

OUTORGADO(S): {outorgado_nome}, advogado(a), inscrito(a) na {outorgado_oab}, com endereço profissional sito à {outorgado_endereco}.

PODERES: Se o tipo for "cível", use exatamente: "Os da Cláusula AD JUDICIA e os especiais para: desistir, transigir, firmar compromissos, assinar documentos, podendo receber notificações, citações e intimações, carta precatória na forma da lei, levantamento de alvará judicial, e poderes especiais para representar o outorgante perante qualquer órgão da administração pública ou privada, praticar todos os demais atos necessários ao fiel desempenho do presente mandato, inclusive substabelecer, com ou sem reservas de poderes."
Se o tipo for "criminal", use o mesmo texto acima e adicione ao final: "nos autos do processo nº {numero_processo}."

{cidade}, _____ de ________________ de 20____.

__________________________________________________
{outorgante_nome}

NÃO adicione nenhum outro campo, assinatura ou texto além do que está neste modelo.""",

    "doc_termo": """Gere um Termo de Declaração de Inaptidão Provisória em português brasileiro seguindo EXATAMENTE esta estrutura, sem mover nenhum elemento de lugar:

{declarante_nome}, {declarante_qualificacao}, inscrito no CPF nº {declarante_cpf}, residente e domiciliado em {declarante_endereco}. Declara para os devidos fins, de fato e de direito, nos termos da lei n.º 7.115/83, art. 2º, parágrafo único da lei n.º 1.060/50 art.98/99 do Novo CPC e art. 5º, inciso LXXIV da Constituição Federal de 1988, que dispõe sobre prova documental para todos os fins de direito, inclusive para fazer prova junto à Justiça Gratuita, que é pobre na forma da lei, não podendo custear as despesas com processo judicial oneroso sem ameaçar a subsistência própria e de sua família, pelo que assume inteira responsabilidade, sob as penas da lei por esta declaração.

{cidade}, ____ de _________ de 20____.

_______________________________________________
{declarante_nome}

REGRAS CRÍTICAS:
- A data ({cidade}, ____ de ___ de 20____) vem DEPOIS do parágrafo da declaração, NUNCA antes
- NÃO adicione nenhum outro campo, assinatura ou texto
- NÃO mova nenhum elemento de lugar
- Copie a estrutura EXATAMENTE como está acima""",

    "doc_contrato": """Gere um Contrato de Honorários Advocatícios em português brasileiro seguindo EXATAMENTE esta estrutura, substituindo apenas os dados variáveis:

Pelo presente instrumento de contrato de prestação de serviços advocatícios, constitui a parte contratante seus bastantes procuradores legais: {advogado_nome}, advogado(a), inscrito(a) na {advogado_oab}, com endereço profissional sito à {advogado_endereco}, doravante denominado CONSTITUIDOS, e do outro lado:

{cliente_nome}, {cliente_qualificacao}, portador(a) do CPF nº {cliente_cpf}, residente e domiciliado em {cliente_endereco}.

I- O constituído se compromete a patrocinar a causa do (a) constituinte, que consiste em ajuizar {objeto}, exercendo todos os atos necessários para a defesa de seus interesses até a prolação da Sentença.

II - Em contraprestação aos serviços prestados, o constituinte se compromete a remunerar os serviços prestados pelo constituído, sob a importância de:
- O contratante pagará a quantia de {valor_entrada} a título de entrada, no ato da assinatura deste contrato.
- Fica estabelecido {percentual} do valor proveito que o contratante obtiver com a ação judicial.

III- O constituinte declara aceitar as condições de caracterizar a presente prestação, independendo, pois, de sucesso na causa, não obstante responda pelas perdas e danos, oriundas de falta de diligência na condução da causa;

IV- A verba oriunda das partes adversas pelo PRINCIPIO DA SUCUMBÊNCIA, REVERTERÁ em beneficio exclusivo do CONSTITUIDO, em limite arbitrado pelo juiz, desvinculado do presente contrato e isento de qualquer desconto;

V - O presente contrato abrange apenas a prestação contida no item I, do presente instrumento, qualquer ação subsequente embora correlata, será objeto de um novo contrato;

VI - As custas, bem como todas as despesas processuais, quer sejam judiciais ou extrajudiciais, inerentes ao presente processo, correrão por conta exclusiva do CONSTITUINTE;

VII - A totalidade dos honorários poderá ser exigida imediatamente, caso haja composição amigável, realizada por quaisquer das partes litigantes, ou em caso de não prosseguimento da ação, por qualquer circunstância, não determinada pelo advogado constituído ou, ainda, se lhe for cassado o mandado sem culpa;

VIII - O presente contrato poderá ser rescindido por qualquer uma das partes sem o pagamento de multa ressaltado o respeito das seguintes providências.
a - Se a rescisão partir do constituído, esta deverá notificar de sua renúncia e aguardar o prazo de 10 dias para a nomeação e um substituto, sem a necessidade de devolução de honorários recebidos, mas desistindo das parcelas futuras;
b - Se a rescisão partir do constituinte, este deverá estar com os honorários devidamente quitados até o momento da rescisão.

IX- Fica eleito o fórum da Comarca de União dos Palmares para dirimir eventuais litígios decorrentes do presente contrato.

Estando as partes de comum acordo, firmam o presente contrato em 02 vias para maior validade jurídica.

{cidade}, _____ de ________________ de 20____.

__________________________________________________
{cliente_nome}

_______________________________________________
{advogado_nome}
{advogado_oab}

Testemunha I ____________________________________________________

Testemunha II ___________________________________________________

NÃO adicione nenhum outro campo ou texto além do que está neste modelo.""",
}


def criar_docx(tipo: str, conteudo: str, dados: dict) -> str:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(TIPOS_DOCUMENTO[tipo].upper())
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Conteúdo gerado pela IA
    for linha in conteudo.split("\n"):
        if linha.strip():
            p = doc.add_paragraph(linha)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(12)
        else:
            doc.add_paragraph()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


async def handle_documento_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("📜 Procuração", callback_data="doc_procuracao")],
        [InlineKeyboardButton("📋 Termo de Declaração", callback_data="doc_termo")],
        [InlineKeyboardButton("📃 Contrato", callback_data="doc_contrato")],
        [InlineKeyboardButton("🔙 Voltar", callback_data="menu_voltar")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if update.callback_query:
        await update.callback_query.edit_message_text(
            "📝 *Geração de Documentos*\n\nQual documento deseja gerar?",
            parse_mode="Markdown",
            reply_markup=reply_markup
        )
    else:
        await update.message.reply_text(
            "📝 *Geração de Documentos*\n\nQual documento deseja gerar?",
            parse_mode="Markdown",
            reply_markup=reply_markup
        )


async def handle_documento_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    doc_tipo = query.data

    if doc_tipo == "menu_voltar":
        keyboard = [
            [InlineKeyboardButton("💬 ChatBot Jurídico", callback_data="menu_chat")],
            [InlineKeyboardButton("📄 Resumir Documento", callback_data="menu_resumo")],
            [InlineKeyboardButton("📝 Gerar Documento", callback_data="menu_documento")],
        ]
        await query.edit_message_text(
            "⚖️ *Menu Principal*\n\nEscolha uma opção:",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return

    if doc_tipo not in CAMPOS_DOCUMENTO:
        return

    context.user_data["mode"] = doc_tipo
    context.user_data["doc_campos"] = CAMPOS_DOCUMENTO[doc_tipo].copy()
    context.user_data["doc_dados"] = {}
    context.user_data["doc_tipo"] = doc_tipo

    campo_key, campo_label = context.user_data["doc_campos"][0]
    context.user_data["doc_campo_atual"] = 0

    await query.edit_message_text(
        f"📝 *{TIPOS_DOCUMENTO[doc_tipo]}*\n\n"
        f"Vou precisar de algumas informações. Você pode digitar /menu a qualquer momento para cancelar.\n\n"
        f"1/{len(context.user_data['doc_campos'])} — {campo_label}",
        parse_mode="Markdown"
    )


async def handle_documento_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc_tipo = context.user_data.get("doc_tipo")
    campos = context.user_data.get("doc_campos", [])
    dados = context.user_data.get("doc_dados", {})
    idx = context.user_data.get("doc_campo_atual", 0)

    if idx >= len(campos):
        return

    campo_key, _ = campos[idx]
    dados[campo_key] = update.message.text
    context.user_data["doc_dados"] = dados

    idx += 1
    context.user_data["doc_campo_atual"] = idx

    if idx < len(campos):
        _, proximo_label = campos[idx]
        await update.message.reply_text(
            f"{idx + 1}/{len(campos)} — {proximo_label}",
            parse_mode="Markdown"
        )
    else:
        # Todos os campos coletados — gera o documento
        await update.message.reply_text("⏳ Gerando documento, aguarde...")
        await update.message.reply_chat_action("typing")

        try:
            prompt = PROMPTS_DOCUMENTO[doc_tipo].format(**dados)
            response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            conteudo = response.choices[0].message.content
            usage = response.usage
            logger.info(
                f"[TOKENS/DOC] input={usage.prompt_tokens} | "
                f"output={usage.completion_tokens} | "
                f"total={usage.total_tokens}"
            )

            docx_path = criar_docx(doc_tipo, conteudo, dados)

            nome_arquivo = f"{TIPOS_DOCUMENTO[doc_tipo].replace(' ', '_')}.docx"
            with open(docx_path, "rb") as f:
                await update.message.reply_document(
                    document=f,
                    filename=nome_arquivo,
                    caption=f"✅ *{TIPOS_DOCUMENTO[doc_tipo]}* gerado com sucesso!\n\n"
                            "⚠️ Revise o documento antes de assinar. Este é um modelo gerado por IA.",
                    parse_mode="Markdown"
                )

            os.unlink(docx_path)

        except Exception as e:
            await update.message.reply_text(
                "⚠️ Erro ao gerar o documento. Tente novamente."
            )

        finally:
            context.user_data.clear()
